import pypandoc
import os
import sys
from docx import Document # python-docx
import subprocess # subprocessモジュールをインポート
import comtypes.client # comtypesを追加
# import comtypes.os_specific # LO_PATHの解決に使う可能性

# LibreOfficeの実行ファイルパス (環境に合わせて調整が必要な場合がある)
# 通常は comtypes が自動で見つけてくれることを期待するが、見つからない場合は指定する
# LO_PATH = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"

def get_libreoffice_path():
    """LibreOfficeの soffice.exe のパスを取得しようと試みる。"""
    # 環境変数から試す
    lo_path_env = os.environ.get('LIBREOFFICE_PROGRAM_PATH')
    if lo_path_env and os.path.exists(os.path.join(lo_path_env, 'soffice.exe')):
        return os.path.join(lo_path_env, 'soffice.exe')
    
    # 一般的なインストール場所を試す (Windows)
    common_paths = [
        'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
        'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',
    ]
    for path in common_paths:
        if os.path.exists(path):
            return path
    
    # comtypes.os_specific から取得を試みる (Windowsのみ) - コメントアウト
    # try:
    #     if hasattr(comtypes.os_specific, 'get_classes_in_module'): # Windowsでのみ利用可能
    #         # この方法は直接的なパス取得ではなく、COMオブジェクトの登録を探すものに近い
    #         # 実際のパスを返すわけではないので、ここでは使用が難しい
    #         pass
    # except Exception:
    #     pass

    return None # 見つからなければNoneを返す

def convert_doc_to_docx_comtypes(input_doc_path, output_docx_path):
    """
    comtypesを使用してLibreOffice API経由で.docファイルを.docxファイルに変換します。
    """
    try:
        abs_input_doc_path = os.path.abspath(input_doc_path)
        abs_output_docx_path = os.path.abspath(output_docx_path)

        print(f"DEBUG: comtypes 入力: {abs_input_doc_path}")
        print(f"DEBUG: comtypes 出力: {abs_output_docx_path}")

        # 出力ディレクトリ作成
        output_dir = os.path.dirname(abs_output_docx_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # LibreOfficeのCOMオブジェクトを取得/起動
        try:
            desktop = comtypes.client.CreateObject("com.sun.star.frame.Desktop")
        except OSError as e:
            print(f"エラー: LibreOfficeのCOMオブジェクトの作成に失敗しました。LibreOfficeが正しくインストールされ、COMコンポーネントが登録されているか確認してください。", file=sys.stderr)
            print(f"詳細: {e}", file=sys.stderr)
            # ここでLibreOfficeの起動を試みることもできるが、より複雑になる
            # lo_path = get_libreoffice_path()
            # if lo_path:
            #     print(f"DEBUG: LibreOfficeの起動を試みます: {lo_path}")
            #     subprocess.Popen([lo_path, '--accept=socket,host=localhost,port=2002;urp;', '--norestore', '--invisible'])
            #     import time
            #     time.sleep(10) # 起動待機（時間は調整が必要）
            #     try:
            #         desktop = comtypes.client.CreateObject("com.sun.star.frame.Desktop")
            #     except Exception as e2:
            #         print(f"エラー: LibreOffice起動後のCOMオブジェクト作成にも失敗しました: {e2}", file=sys.stderr)
            #         return False
            # else:
            #     print(f"エラー: LibreOfficeの実行パスが見つからず、COMオブジェクトも作成できませんでした。", file=sys.stderr)
            #     return False
            return False # 現状は起動試行はコメントアウト

        # ドキュメントを開くためのプロパティ
        # file:/// URI形式でパスを指定するのが確実
        input_url = uno_path(abs_input_doc_path)
        # 読み込み専用、非表示で開く
        props = (
            comtypes.client.PropertyValue(Name="ReadOnly", Value=True),
            comtypes.client.PropertyValue(Name="Hidden", Value=True)
        )
        
        doc = desktop.loadComponentFromURL(input_url, "_blank", 0, props)
        if not doc:
            print(f"エラー: LibreOfficeでドキュメントを開けませんでした: {abs_input_doc_path}", file=sys.stderr)
            return False

        # 保存するためのプロパティ
        output_url = uno_path(abs_output_docx_path)
        save_props = (
            comtypes.client.PropertyValue(Name="FilterName", Value="Office Open XML Text"), # MS Word 2007-2013 XML (*.docx)
            # または "MS Word 2007 XML" など、LibreOfficeのバージョンや設定によって名前が違う可能性あり。
            # 正確なフィルター名は、LibreOfficeのマクロで確認するか、ドキュメントを参照。
            comtypes.client.PropertyValue(Name="Overwrite", Value=True)
        )

        doc.storeToURL(output_url, save_props)
        print(f"DOCからDOCXへの変換成功(comtypes): {abs_input_doc_path} -> {abs_output_docx_path}")
        
        # ドキュメントを閉じる
        if hasattr(doc, 'close'):
            doc.close(True)
        
        return True

    except Exception as e:
        print(f"エラー: DOCからDOCXへの変換中に予期せぬエラーが発生しました (comtypes): {type(e).__name__} - {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return False

# comtypes.client.CreateObject のために必要になることがあるヘルパークラス
class PropertyValue:
    def __init__(self, Name, Value):
        self.Name = Name
        self.Value = Value

# PythonパスをUNOのfile URLに変換するヘルパー関数
def uno_path(path):
    if sys.platform == 'win32':
        return "file:///" + path.replace('\\', '/').replace(':', '|') # Windows形式
    else:
        return "file:///" + path # Linux/Mac形式 (要検証)

# subprocess を使用する古い関数 (UnicodeDecodeError抑制済み)
def convert_doc_to_docx_libreoffice(input_doc_path, output_docx_path):
    """
    LibreOfficeをsubprocess経由で呼び出し、.docファイルを.docxファイルに変換します。
    出力ファイルは input_doc_path と同じディレクトリに、input_doc_path のベース名 + .docx として生成され、
    その後、指定された output_docx_path にリネーム（または既に同じなら何もしない）されます。
    [注意] この関数は日本語ファイル名の扱いに問題を抱えている可能性があります。

    :param input_doc_path: 入力.docファイルのパス
    :param output_docx_path: 目的の出力.docxファイルのパス
    :return: 変換が成功した場合はTrue、失敗した場合はFalse
    """
    try:
        abs_input_doc_path = os.path.abspath(input_doc_path)
        abs_output_docx_path = os.path.abspath(output_docx_path)

        output_dir_for_libreoffice = os.path.dirname(abs_input_doc_path) 
        
        temp_docx_filename_by_libreoffice = os.path.splitext(os.path.basename(abs_input_doc_path))[0] + ".docx"
        temp_abs_output_docx_path_by_libreoffice = os.path.join(output_dir_for_libreoffice, temp_docx_filename_by_libreoffice)

        print(f"DEBUG: LibreOffice (subprocess) 入力: {abs_input_doc_path}")
        print(f"DEBUG: LibreOffice (subprocess) 一時出力先ディレクトリ: {output_dir_for_libreoffice}")
        print(f"DEBUG: LibreOffice (subprocess) が生成する一時ファイル名: {temp_abs_output_docx_path_by_libreoffice}")
        print(f"DEBUG: LibreOffice (subprocess) 最終的なDOCX出力パス: {abs_output_docx_path}")

        final_output_dir = os.path.dirname(abs_output_docx_path)
        if final_output_dir and not os.path.exists(final_output_dir):
            os.makedirs(final_output_dir)

        command = [
            'C:\\\\Program Files\\\\LibreOffice\\\\program\\\\soffice.exe', # フルパス指定 (.exe)
            '--headless',
            '--invisible',
            '--convert-to', 'docx',
            '--outdir', output_dir_for_libreoffice, 
            abs_input_doc_path
        ]
        
        print(f"DEBUG: LibreOfficeコマンド (subprocess): {' '.join(command)}")
        process = subprocess.run(command, capture_output=True, text=True, encoding='utf-8', check=False, errors='ignore')

        if process.returncode == 0:
            print(f"DEBUG: LibreOffice (subprocess) stdout: {process.stdout}")
            if process.stderr:
                print(f"DEBUG: LibreOffice (subprocess) stderr: {process.stderr}")

            if os.path.exists(temp_abs_output_docx_path_by_libreoffice):
                print(f"LibreOffice (subprocess) による一時ファイル生成成功: {temp_abs_output_docx_path_by_libreoffice}")
                if temp_abs_output_docx_path_by_libreoffice != abs_output_docx_path:
                    try:
                        os.rename(temp_abs_output_docx_path_by_libreoffice, abs_output_docx_path)
                        print(f"DOCからDOCXへの変換成功(subprocess): {abs_input_doc_path} -> {abs_output_docx_path}")
                        return True
                    except OSError as e:
                        print(f"エラー: 一時ファイルのリネーム/移動に失敗しました (subprocess): {e}", file=sys.stderr)
                        if os.path.exists(temp_abs_output_docx_path_by_libreoffice):
                            try:
                                os.remove(temp_abs_output_docx_path_by_libreoffice)
                            except OSError: 
                                pass
                        return False
                else: 
                    print(f"DOCからDOCXへの変換成功(subprocess): {abs_input_doc_path} -> {abs_output_docx_path}")
                    return True
            else:
                print(f"エラー: LibreOfficeコマンド成功しましたが、一時出力ファイルが見つかりません (subprocess): {temp_abs_output_docx_path_by_libreoffice}", file=sys.stderr)
                return False
        else:
            print(f"エラー: LibreOfficeコマンド実行失敗 (subprocess) (終了コード: {process.returncode})", file=sys.stderr)
            return False

    except FileNotFoundError: 
        print(f"エラー: LibreOfficeコマンド(soffice)が見つかりません (subprocess)。LibreOfficeがインストールされ、PATHが通っているか確認してください。", file=sys.stderr)
        return False
    except Exception as e:
        print(f"エラー: DOCからDOCXへの変換中に予期せぬエラーが発生しました (subprocess LibreOffice): {type(e).__name__} - {e}", file=sys.stderr)
        return False

def convert_docx_to_txt(input_docx_path, output_txt_path):
    """
    python-docxを使用して.docxファイルからテキストを抽出し、.txtファイルとして保存します。

    :param input_docx_path: 入力.docxファイルのパス
    :param output_txt_path: 出力.txtファイルのパス
    :return: 変換が成功した場合はTrue、失敗した場合はFalse
    """
    try:
        abs_input_docx_path = os.path.abspath(input_docx_path)
        abs_output_txt_path = os.path.abspath(output_txt_path)
        print(f"DEBUG: python-docx 入力: {abs_input_docx_path}")
        print(f"DEBUG: python-docx 出力: {abs_output_txt_path}")

        if not os.path.exists(abs_input_docx_path):
            print(f"エラー: python-docxの入力ファイルが見つかりません: {abs_input_docx_path}", file=sys.stderr)
            return False

        doc = Document(abs_input_docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        output_dir = os.path.dirname(abs_output_txt_path)
        if output_dir:
            if not os.path.exists(output_dir):
                print(f"DEBUG: 出力ディレクトリを作成します: {output_dir}")
                os.makedirs(output_dir)
            else:
                print(f"DEBUG: 出力ディレクトリは既に存在します: {output_dir}")
        else:
            print(f"DEBUG: 出力ディレクトリは指定されていません (カレントディレクトリ想定)")
            
        with open(abs_output_txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
        
        if os.path.exists(abs_output_txt_path):
            print(f"DOCXからTXTへの変換成功(ファイル確認OK): {abs_input_docx_path} -> {abs_output_txt_path}")
            return True
        else:
            print(f"エラー: TXTファイル書き込み成功のはずが出力ファイルが見つかりません: {abs_output_txt_path}", file=sys.stderr)
            return False
            
    except FileNotFoundError: # Document()で発生する可能性
        print(f"エラー: 入力DOCXファイルが見つかりません(python-docx): {abs_input_docx_path}", file=sys.stderr)
        return False
    except Exception as e:
        print(f"エラー: DOCXからTXTへの変換中に予期せぬエラーが発生しました: {type(e).__name__} - {e}", file=sys.stderr)
        return False

if __name__ == '__main__':
    if len(sys.argv) == 3:
        input_doc_file = sys.argv[1]
        output_txt_file = sys.argv[2]

        if not input_doc_file.lower().endswith(".doc"):
            print("エラー: 入力ファイルは.docである必要があります。", file=sys.stderr)
            sys.exit(1)
        if not output_txt_file.lower().endswith(".txt"):
            print("エラー: 出力ファイルは.txtである必要があります。", file=sys.stderr)
            sys.exit(1)

        base_name = os.path.splitext(input_doc_file)[0]
        intermediate_docx_file = base_name + ".docx"
        
        print(f"DEBUG: メイン処理 - 入力DOC: {input_doc_file}")
        print(f"DEBUG: メイン処理 - 中間DOCX: {intermediate_docx_file}")
        print(f"DEBUG: メイン処理 - 出力TXT: {output_txt_file}")

        # .docから.docxへの変換 (comtypes版を使用)
        if convert_doc_to_docx_comtypes(input_doc_file, intermediate_docx_file):
            if not convert_docx_to_txt(intermediate_docx_file, output_txt_file):
                print(f"エラー: {intermediate_docx_file} から {output_txt_file} への変換に失敗しました。", file=sys.stderr)
                sys.exit(1)
        else:
            print(f"エラー: {input_doc_file} から {intermediate_docx_file} への変換に失敗しました。comtypes版のログを確認してください。", file=sys.stderr)
            # フォールバックとしてsubprocess版を試すこともできるが、今回はエラーで終了
            # print(f"DEBUG: comtypes版での変換に失敗したため、subprocess版を試行します。")
            # if convert_doc_to_docx_libreoffice(input_doc_file, intermediate_docx_file):
            #     if not convert_docx_to_txt(intermediate_docx_file, output_txt_file):
            #         print(f"エラー: {intermediate_docx_file} から {output_txt_file} への変換に失敗しました。(subprocessフォールバック後)", file=sys.stderr)
            #         sys.exit(1)
            # else:
            #     print(f"エラー: {input_doc_file} から {intermediate_docx_file} への変換に失敗しました。(subprocessフォールバックも失敗)", file=sys.stderr)
            #     sys.exit(1)
            sys.exit(1)
            
    else:
        print("使用方法: python word-to-txt.py <input_doc_file> <output_txt_file>")
        print("例: python word-to-txt.py example.doc output.txt")
