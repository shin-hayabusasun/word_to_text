#!/usr/bin/env python
# coding: utf-8

import os
import sys
import subprocess
import docx2txt
import docx
from pathlib import Path
import re

def extract_text_from_binary(file_path, encoding='utf-8'):
    """
    バイナリファイルから直接テキストを抽出する
    """
    # バイナリモードでファイルを開く
    with open(file_path, 'rb') as f:
        content = f.read()
    
    # 複数のエンコーディングで試す
    encodings = ['utf-8', 'shift_jis', 'euc_jp', 'cp932']
    best_text = None
    best_jp_ratio = 0
    
    for enc in encodings:
        try:
            text = content.decode(enc, errors='ignore')
            
            # 日本語文字の割合を計算
            jp_chars = len([c for c in text if '\u3040' <= c <= '\u30ff' or '\u4e00' <= c <= '\u9fff'])
            if len(text) > 0:
                jp_ratio = jp_chars / len(text)
            else:
                jp_ratio = 0
            
            print(f"  {enc}: 日本語文字比率 {jp_ratio:.2%}")
            
            if jp_ratio > best_jp_ratio:
                best_text = text
                best_jp_ratio = jp_ratio
        except Exception:
            continue
    
    # 十分な日本語文字がある場合のみ
    if best_text and best_jp_ratio > 0.001:  # 0.1%以上が日本語文字
        # テキストをクリーンアップ（余分な制御文字や記号を削除）
        cleaned_text = re.sub(r'[^\x20-\x7E\u3000-\u30FF\u4E00-\u9FFF\u3040-\u309F\uFF00-\uFF9F\u2000-\u206F\n]+', '', best_text)
        return cleaned_text
    
    return None

def convert_doc_or_docx(file_path, output_path=None, encoding='utf-8'):
    """
    .doc/.docxファイルをテキストに変換する
    """
    try:
        # 出力パスが指定されていない場合は入力ファイルと同じ場所に.txtを作成
        if output_path is None:
            output_path = str(Path(file_path).with_suffix('.txt'))
        
        print(f"変換中: {file_path}")
        print(f"出力先: {output_path}")
        
        # ファイル拡張子を確認
        is_docx = file_path.lower().endswith('.docx')
        
        # ファイルの処理
        if is_docx:
            # .docxファイルの処理
            try:
                text = docx2txt.process(file_path)
                with open(output_path, 'w', encoding=encoding) as f:
                    f.write(text)
                print("docx2txtで変換成功")
                return True
            except Exception as e:
                print(f"docx2txtでの変換に失敗: {str(e)}")
                
                try:
                    doc = docx.Document(file_path)
                    full_text = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)
                    
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text)
                            if row_text:
                                full_text.append('\t'.join(row_text))
                    
                    with open(output_path, 'w', encoding=encoding) as f:
                        f.write('\n'.join(full_text))
                    
                    print("python-docxで変換成功")
                    return True
                except Exception as e2:
                    print(f"python-docxでの変換に失敗: {str(e2)}")
        
        # .docファイルまたは.docxの処理が失敗した場合、バイナリからの直接抽出を試みる
        print("バイナリ解析による変換を試みます...")
        text = extract_text_from_binary(file_path)
        
        if text:
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(text)
            print("バイナリ解析による変換成功")
            return True
        else:
            print("バイナリ解析でも有効なテキストが抽出できませんでした")
        
        return False
    
    except Exception as e:
        print(f"変換エラー: {str(e)}")
        return False

def main():
    if len(sys.argv) < 2:
        print("使用方法: python word_converter.py <Word文書のパス> [--encoding=ENCODING]")
        print("  --encoding=ENCODING: 出力エンコーディング (デフォルト: utf-8, 例: --encoding=shift-jis)")
        return
    
    file_path = sys.argv[1]
    encoding = 'utf-8'
    
    # コマンドライン引数を解析
    for arg in sys.argv[2:]:
        if arg.startswith("--encoding="):
            encoding = arg.split("=")[1]
    
    if not os.path.exists(file_path):
        print(f"エラー: 指定されたファイル '{file_path}' が存在しません。")
        return
    
    # ファイル拡張子をチェック
    if not (file_path.lower().endswith('.doc') or file_path.lower().endswith('.docx')):
        print(f"エラー: 指定されたファイル '{file_path}' はWord文書(.docまたは.docx)ではありません。")
        return
    
    # ファイルを変換
    if convert_doc_or_docx(file_path, encoding=encoding):
        print(f"変換完了: {file_path}")
    else:
        print(f"変換失敗: {file_path}")

if __name__ == "__main__":
    main() 