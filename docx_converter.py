#!/usr/bin/env python
# coding: utf-8

import os
import sys
import docx2txt
import docx
from pathlib import Path

def convert_docx_file(docx_path, output_path=None, encoding='utf-8'):
    """
    .docxファイルをテキストファイルに変換する
    """
    try:
        # 出力パスが指定されていない場合は入力ファイルと同じ場所に.txtを作成
        if output_path is None:
            output_path = str(Path(docx_path).with_suffix('.txt'))
        
        print(f"変換中: {docx_path}")
        print(f"出力先: {output_path}")
        
        # docx2txtを使用してテキスト抽出を試みる
        try:
            text = docx2txt.process(docx_path)
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(text)
            print(f"docx2txtで変換成功")
            return True
        except Exception as e:
            print(f"docx2txtでの変換に失敗しました: {str(e)}")
            print("python-docxでの変換を試みます...")
        
        # python-docxを使用して抽出を試みる
        doc = docx.Document(docx_path)
        
        # テキストを抽出
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # 空の段落を無視
                full_text.append(para.text)
        
        # テーブルからもテキストを抽出
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # 空のセルを無視
                        row_text.append(cell.text)
                if row_text:  # 空の行を無視
                    full_text.append('\t'.join(row_text))
        
        # テキストファイルに書き込む
        with open(output_path, 'w', encoding=encoding) as f:
            f.write('\n'.join(full_text))
        
        print(f"python-docxで変換成功")
        return True
    
    except Exception as e:
        print(f"変換エラー: {str(e)}")
        return False

def main():
    if len(sys.argv) < 2:
        print("使用方法: python docx_converter.py <docxファイル> [--encoding=ENCODING]")
        print("  --encoding=ENCODING: 出力エンコーディング (デフォルト: utf-8, 例: --encoding=shift-jis)")
        return
    
    docx_path = sys.argv[1]
    encoding = 'utf-8'
    
    # コマンドライン引数を解析
    for arg in sys.argv[2:]:
        if arg.startswith("--encoding="):
            encoding = arg.split("=")[1]
    
    if not os.path.exists(docx_path):
        print(f"エラー: 指定されたファイル '{docx_path}' が存在しません。")
        return
    
    if not docx_path.lower().endswith('.docx'):
        print(f"エラー: 指定されたファイル '{docx_path}' は.docx形式ではありません。")
        return
    
    # ファイルを変換
    if convert_docx_file(docx_path, encoding=encoding):
        print(f"変換完了: {docx_path}")
    else:
        print(f"変換失敗: {docx_path}")

if __name__ == "__main__":
    main() 