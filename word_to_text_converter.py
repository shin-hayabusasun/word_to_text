#!/usr/bin/env python
# coding: utf-8

import os
import sys
import glob
import docx
import win32com.client
import tempfile
import shutil
import subprocess
import time
import platform
import re
from pathlib import Path
import codecs
import docx2txt
import struct
import binascii
import traceback

def convert_docx_to_text(docx_path, output_path=None):
    """
    .docxファイルをテキストファイルに変換する
    
    Args:
        docx_path (str): 変換するdocxファイルのパス
        output_path (str, optional): 出力先のパス。指定がない場合は同じ場所に.txtファイルを作成
    
    Returns:
        str: 作成されたテキストファイルのパス
    """
    try:
        # 出力パスが指定されていない場合は入力ファイルと同じ場所に.txtを作成
        if output_path is None:
            output_path = str(Path(docx_path).with_suffix('.txt'))
        
        # docx2txtを使用してテキスト抽出を試みる
        try:
            text = docx2txt.process(docx_path)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return output_path
        except Exception as e1:
            print(f"docx2txtでの変換に失敗しました（{docx_path}）: {str(e1)}")
            print("python-docxでの変換を試みます...")
        
        # python-docxを使用して抽出を試みる
        doc = docx.Document(docx_path)
        
        # テキストを抽出
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # 空の段落を無視
                full_text.append(para.text)
        
        # テーブルからもテキストを抽出
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # 空のセルを無視
                        row_text.append(cell.text)
                if row_text:  # 空の行を無視
                    full_text.append('\t'.join(row_text))
        
        # テキストファイルに書き込む
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
        
        return output_path
    
    except PermissionError:
        print(f"変換エラー（{docx_path}）: ファイルにアクセスする権限がありません。ファイルが開かれていないか確認してください。")
        return None
    except docx.opc.exceptions.PackageNotFoundError:
        print(f"変換エラー（{docx_path}）: ファイルが見つからないか、正しいWord文書形式ではありません。")
        return None
    except Exception as e:
        print(f"変換エラー（{docx_path}）: {str(e)}")
        return None

def convert_doc_to_text(doc_path, output_path=None):
    """
    古い形式の.docファイルをテキストファイルに変換する
    
    Args:
        doc_path (str): 変換するdocファイルのパス
        output_path (str, optional): 出力先のパス。指定がない場合は同じ場所に.txtファイルを作成
    
    Returns:
        str: 作成されたテキストファイルのパス
    """
    try:
        # 出力パスが指定されていない場合は入力ファイルと同じ場所に.txtを作成
        if output_path is None:
            output_path = str(Path(doc_path).with_suffix('.txt'))
        
        # 複数の変換方法を順番に試す（優先度順）
        methods = [
            (extract_text_with_word_com_direct, "Word COMでの直接抽出"),
            (extract_japanese_text_enhanced, "強化版日本語特化処理"),
            (extract_text_doc_to_docx, "docからdocxへの変換を経由"),
            (extract_text_with_antiword, "antiwordを使用"),
            (extract_text_with_binary_parsing, "バイナリ解析")
        ]
        
        all_extracted_texts = []
        for extract_func, method_name in methods:
            try:
                print(f"{method_name}での変換を試みます（{doc_path}）...")
                temp_output = f"{output_path}.temp_{method_name.replace(' ', '_')}.txt"
                
                # 抽出関数を実行
                extract_func(doc_path, temp_output)
                
                # 結果を確認
                if os.path.exists(temp_output):
                    with open(temp_output, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    
                    # 有効な日本語テキストかどうかをチェック
                    jp_chars = re.findall(r'[ぁ-んァ-ヶ一-龠々〆〜]', text)
                    jp_ratio = len(jp_chars) / max(len(text), 1)
                    
                    print(f"  {method_name}: テキスト長={len(text)}, 日本語文字数={len(jp_chars)}, 比率={jp_ratio:.2%}")
                    
                    # 十分な長さと日本語比率があれば保存
                    if len(text) > 100 and jp_ratio > 0.05:
                        all_extracted_texts.append((text, jp_ratio, method_name, temp_output))
                    else:
                        print(f"  {method_name}: 十分な日本語テキストが含まれていません")
                        os.remove(temp_output)
                else:
                    print(f"  {method_name}: 出力ファイルが生成されませんでした")
            except Exception as e:
                print(f"  {method_name}での変換に失敗: {str(e)}")
        
        # 結果を評価して最適なものを選択
        if all_extracted_texts:
            # 日本語比率とテキスト長で並べ替え
            all_extracted_texts.sort(key=lambda x: (x[1], len(x[0])), reverse=True)
            best_text, best_ratio, best_method, best_temp = all_extracted_texts[0]
            
            print(f"最適な変換結果: {best_method} (日本語比率: {best_ratio:.2%}, 文字数: {len(best_text)})")
            
            # テキストの後処理
            processed_text = best_text
            # 余分なマーカーなどを削除
            processed_text = re.sub(r'^---.*?---\n', '', processed_text)
            processed_text = re.sub(r'\n---.*?---\n', '\n', processed_text)
            # 制御文字や特殊記号を削除
            processed_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', processed_text)
            # 連続した空行を整理
            processed_text = re.sub(r'\n{3,}', '\n\n', processed_text)
            
            # 最終テキストを出力
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(processed_text)
            
            # 一時ファイルを削除
            for _, _, _, temp_file in all_extracted_texts:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except:
                    pass
            
            print(f"変換完了: {output_path}")
            return output_path
        else:
            # すべての方法が失敗した場合は最終手段としてバイナリデータから直接抽出
            print("すべての方法が失敗したため、バイナリデータから直接抽出します...")
            return extract_japanese_text_enhanced(doc_path, output_path)
    
    except Exception as e:
        print(f"変換エラー: {str(e)}")
        traceback.print_exc()
        raise

def extract_japanese_text_enhanced(doc_path, output_path):
    """
    日本語テキスト抽出に特化した強化版処理
    このメソッドは特にWordバイナリファイル内の日本語テキストの検出と抽出に焦点を当てています
    """
    try:
        # バイナリモードでファイルを開く
        with open(doc_path, 'rb') as f:
            content = f.read()
        
        # 複数のエンコーディングで抽出を試み、最も多くのテキストを取得したものを採用
        all_extracted_texts = []
        
        # 1. Win32 COMによる直接抽出を試みる（Windows環境のみ）
        if platform.system() == 'Windows':
            try:
                import win32com.client
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                doc = word_app.Documents.Open(os.path.abspath(doc_path), ReadOnly=True)
                text = doc.Content.Text
                doc.Close(SaveChanges=False)
                word_app.Quit()
                
                if text and len(text) > 100:  # 十分なテキストが取得できた場合
                    # 日本語文字の比率を計算
                    jp_chars = re.findall(r'[ぁ-んァ-ヶ一-龠々〆〜]', text)
                    jp_ratio = len(jp_chars) / len(text) if len(text) > 0 else 0
                    
                    if jp_ratio > 0.05:  # 5%以上が日本語文字である場合
                        all_extracted_texts.append((text, jp_ratio, "win32com"))
            except Exception as e:
                print(f"COM抽出失敗: {str(e)}")
        
        # 2. Python-docxによる抽出（.docxファイル用）
        try:
            import docx
            doc = docx.Document(doc_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = '\n'.join(paragraphs)
            
            if text and len(text) > 100:
                # 日本語文字の比率を計算
                jp_chars = re.findall(r'[ぁ-んァ-ヶ一-龠々〆〜]', text)
                jp_ratio = len(jp_chars) / len(text) if len(text) > 0 else 0
                
                if jp_ratio > 0.05:
                    all_extracted_texts.append((text, jp_ratio, "python-docx"))
        except Exception as e:
            print(f"python-docx抽出失敗: {str(e)}")
        
        # 3. バイナリ解析による抽出
        # 3.1 UTF-16LE (Windows標準のUnicode形式)での抽出
        try:
            # バイナリデータから直接UTF-16LEの日本語テキストを抽出
            # Word文書ではテキストがUTF-16LEで格納されていることが多い
            utf16_text = ""
            i = 0
            text_chunks = []
            current_chunk = []
            
            while i < len(content) - 2:
                # 2バイト単位で読み取り
                if i + 2 <= len(content):
                    try:
                        char_code = struct.unpack('<H', content[i:i+2])[0]
                        
                        # 以下の範囲は日本語文字の領域
                        is_japanese = (
                            (0x3040 <= char_code <= 0x309F) or  # ひらがな
                            (0x30A0 <= char_code <= 0x30FF) or  # カタカナ
                            (0x4E00 <= char_code <= 0x9FFF) or  # 漢字
                            (char_code in [0x3000, 0x3001, 0x3002, 0xFF01, 0xFF0C, 0xFF0E, 0xFF1A, 0xFF1F, 
                                         0x2025, 0x2026, 0x301C, 0x303B, 0x30FB]) or  # 句読点など
                            (0xFF10 <= char_code <= 0xFF19) or  # 全角数字
                            (0xFF21 <= char_code <= 0xFF3A) or  # 全角英大文字
                            (0xFF41 <= char_code <= 0xFF5A)     # 全角英小文字
                        )
                        
                        if is_japanese:
                            try:
                                char = content[i:i+2].decode('utf-16le', errors='ignore')
                                current_chunk.append(char)
                            except:
                                pass
                        else:
                            # 特定の記号も取り込む（文章の繋がりに重要）
                            is_valid_char = (
                                char_code in [0x0020, 0x0009, 0x000A, 0x000D, 0x3000, 0x00A0, 0x0026] or  # スペース、タブ、改行など
                                (0x0030 <= char_code <= 0x0039) or  # 数字
                                (0x0041 <= char_code <= 0x005A) or  # 英大文字
                                (0x0061 <= char_code <= 0x007A) or  # 英小文字
                                (char_code in [0x002E, 0x002C, 0x003A, 0x003B, 0x0028, 0x0029, 0x005B, 0x005D, 
                                            0x007B, 0x007D, 0x0025, 0x002F, 0x2019, 0x2026])  # 記号
                            )
                            
                            if is_valid_char:
                                try:
                                    char = content[i:i+2].decode('utf-16le', errors='ignore')
                                    current_chunk.append(char)
                                except:
                                    pass
                            # 無関係な文字が一定数連続したら新しいチャンクとする
                            elif len(current_chunk) > 15:  # 充分な長さのチャンクが溜まった
                                chunk_text = ''.join(current_chunk)
                                # 日本語テキストが含まれるチャンクのみ保存
                                if re.search(r'[ぁ-んァ-ヶ一-龠々〆〜]', chunk_text):
                                    text_chunks.append(chunk_text)
                                current_chunk = []
                            else:
                                # チャンクが短すぎる場合はリセット
                                current_chunk = []
                    except:
                        pass
                i += 2
            
            # 最後のチャンクも追加
            if current_chunk and len(current_chunk) > 15:
                chunk_text = ''.join(current_chunk)
                if re.search(r'[ぁ-んァ-ヶ一-龠々〆〜]', chunk_text):
                    text_chunks.append(chunk_text)
            
            # チャンク間の重複を除去して連結
            if text_chunks:
                unique_chunks = []
                for chunk in text_chunks:
                    if not any(chunk in uc for uc in unique_chunks):
                        unique_chunks.append(chunk)
                
                utf16_text = '\n'.join(unique_chunks)
                
                # 不要なバイナリデータやノイズを除去
                utf16_text = re.sub(r'[^\x20-\x7E\u3000-\u30FF\u4E00-\u9FFF\u3040-\u309F\uFF00-\uFF9F\u2000-\u206F\n]+', '', utf16_text)
                utf16_text = re.sub(r'[\x00-\x1F\x7F]', '', utf16_text)  # 制御文字を除去
                
                # 日本語比率を計算
                jp_chars = re.findall(r'[ぁ-んァ-ヶ一-龠々〆〜]', utf16_text)
                jp_ratio = len(jp_chars) / len(utf16_text) if len(utf16_text) > 0 else 0
                
                if len(utf16_text) > 100 and jp_ratio > 0.1:
                    all_extracted_texts.append((utf16_text, jp_ratio, "binary_utf16"))
        
        except Exception as e:
            print(f"バイナリ解析（UTF-16）失敗: {str(e)}")
        
        # 3.2 その他のエンコーディングでの抽出を試みる
        encodings = ['utf-8', 'shift_jis', 'euc-jp', 'cp932', 'iso-2022-jp']
        
        for encoding in encodings:
            try:
                decoded_text = content.decode(encoding, errors='ignore')
                
                # 日本語文字が含まれているか確認
                jp_chars = re.findall(r'[ぁ-んァ-ヶ一-龠々〆〜]', decoded_text)
                jp_ratio = len(jp_chars) / len(decoded_text) if len(decoded_text) > 0 else 0
                
                if len(decoded_text) > 100 and jp_ratio > 0.05:
                    # 意味のある段落を抽出
                    lines = decoded_text.splitlines()
                    meaningful_lines = []
                    
                    for line in lines:
                        clean_line = re.sub(r'[\x00-\x1F\x7F]', '', line)  # 制御文字を除去
                        if re.search(r'[ぁ-んァ-ヶ一-龠々〆〜]', clean_line) and len(clean_line.strip()) > 3:
                            meaningful_lines.append(clean_line)
                    
                    if meaningful_lines:
                        clean_text = '\n'.join(meaningful_lines)
                        # 再度日本語比率を確認
                        jp_chars = re.findall(r'[ぁ-んァ-ヶ一-龠々〆〜]', clean_text)
                        jp_ratio = len(jp_chars) / len(clean_text) if len(clean_text) > 0 else 0
                        
                        if len(clean_text) > 100 and jp_ratio > 0.1:
                            all_extracted_texts.append((clean_text, jp_ratio, f"encoding_{encoding}"))
            except Exception as e:
                print(f"{encoding}でのデコード失敗: {str(e)}")
        
        # 抽出結果を評価して最適なものを選択
        if all_extracted_texts:
            # 日本語比率とテキスト長で並べ替え（日本語比率を優先）
            sorted_texts = sorted(all_extracted_texts, key=lambda x: (x[1], len(x[0])), reverse=True)
            best_text, best_ratio, best_method = sorted_texts[0]
            
            print(f"最適な抽出方法: {best_method} (日本語比率: {best_ratio:.2%}, 文字数: {len(best_text)})")
            
            # テキストの後処理
            # 不要なラベルを除去
            best_text = re.sub(r'^---.*?---\s*', '', best_text)
            best_text = re.sub(r'\n---.*?---\s*', '\n', best_text)
            
            # 重複行を除去
            lines = best_text.splitlines()
            unique_lines = []
            for line in lines:
                clean_line = line.strip()
                if clean_line and clean_line not in unique_lines:
                    unique_lines.append(clean_line)
            
            # 余分な空行を整理
            consolidated_text = '\n'.join(unique_lines)
            
            # XMLタグや不要なマークアップを除去
            consolidated_text = re.sub(r'<.*?>', '', consolidated_text)
            consolidated_text = re.sub(r'\\+[a-zA-Z]+', '', consolidated_text)
            
            # 「生命保険協会ガイドライン」というキーワードが含まれているかチェック
            if '生命保険協会ガイドライン' in consolidated_text:
                # このキーワードを含む部分からテキストを再構成
                start_idx = consolidated_text.find('生命保険協会ガイドライン')
                if start_idx >= 0:
                    consolidated_text = consolidated_text[start_idx:]
            
            # テキストファイルに書き込む
            with open(output_path, 'w', encoding='utf-8') as out_file:
                out_file.write(consolidated_text)
            
            print(f"日本語テキスト抽出完了: {output_path}")
            return output_path
        else:
            raise Exception("有効な日本語テキストが見つかりませんでした")
    
    except Exception as e:
        print(f"日本語テキスト抽出失敗: {str(e)}")
        traceback.print_exc()
        raise

def extract_text_with_word_com_direct(doc_path, output_path):
    """
    Word COMを使用して直接テキストを抽出する
    """
    # 絶対パスに変換
    doc_path = os.path.abspath(doc_path)
    output_path = os.path.abspath(output_path)
    
    try:
        # Wordアプリケーションの起動
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        
        try:
            # docファイルを開く
            print(f"Word COMでファイルを開いています: {doc_path}")
            doc = word.Documents.Open(doc_path, ReadOnly=True)
            
            # テキストを直接抽出
            print("ドキュメントからテキストを抽出中...")
            text = doc.Content.Text
            
            # docファイルを閉じる
            doc.Close(SaveChanges=False)
            
            # テキストファイルに書き込む
            print(f"テキストをファイルに書き込み中: {output_path}")
            with open(output_path, 'w', encoding='utf-8', errors='ignore') as f:
                f.write(text)
            
            return output_path
        
        except Exception as e:
            print(f"Word COM直接テキスト抽出エラー（{doc_path}）: {str(e)}")
            traceback.print_exc()
            raise e
        
        finally:
            # Wordアプリケーションを終了
            try:
                word.Quit()
            except:
                pass
    
    except Exception as e:
        print(f"Word COM初期化エラー: {str(e)}")
        traceback.print_exc()
        raise e

def extract_text_with_japanese_support(doc_path, output_path):
    """
    日本語テキスト抽出に特化したカスタム処理
    """
    try:
        # バイナリモードでファイルを開く
        with open(doc_path, 'rb') as f:
            content = f.read()
        
        # バイナリからの日本語テキスト抽出
        # 日本語のShift-JIS, EUC-JP, UTF-8で抽出を試みる
        extracted_text = ""
        encodings = ['utf-8', 'shift_jis', 'euc-jp', 'iso-2022-jp', 'cp932']
        
        # バイナリデータから日本語テキスト部分を検出
        for enc in encodings:
            try:
                # バイナリをデコードしてみる
                decoded = content.decode(enc, errors='ignore')
                
                # 日本語文字が含まれているかチェック
                if re.search(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]', decoded):
                    # 日本語文字が含まれている場合
                    extracted_text = decoded
                    break
            except UnicodeDecodeError:
                continue
        
        # 日本語テキストが検出されなかった場合、より高度な処理
        if not extracted_text or not re.search(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]', extracted_text):
            # バイナリデータから日本語文字のシーケンスを検出
            # Word文書の中から日本語テキスト部分を特定する
            text_chunks = []
            i = 0
            while i < len(content) - 1:
                # 2バイト文字の検出（日本語は主に2バイト文字）
                if 0x81 <= content[i] <= 0x9F or 0xE0 <= content[i] <= 0xEF:
                    # Shift-JISの範囲内の場合
                    try:
                        # 2バイト分取得してデコード
                        char_bytes = content[i:i+2]
                        char = char_bytes.decode('shift_jis', errors='ignore')
                        if re.search(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]', char):
                            text_chunks.append(char)
                        i += 2
                    except:
                        i += 1
                else:
                    # ASCII文字の場合
                    if 32 <= content[i] <= 126:  # 表示可能なASCII
                        text_chunks.append(chr(content[i]))
                    i += 1
            
            # 収集した文字列を連結
            extracted_text = ''.join(text_chunks)
            
            # 不要な制御文字を削除
            extracted_text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', extracted_text)
            
            # 連続する空白を1つに
            extracted_text = re.sub(r'\s+', ' ', extracted_text)
        
        # テキストファイルに書き込む
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(extracted_text)
        
        # 書き込まれたファイルが空でないか確認
        if os.path.getsize(output_path) > 0:
            return output_path
        else:
            raise Exception("抽出されたテキストが空です")
    
    except Exception as e:
        print(f"日本語テキスト抽出エラー: {str(e)}")
        raise e

def extract_text_with_antiword(doc_path, output_path):
    """
    antiwordライブラリを使用してdocファイルからテキストを抽出する
    """
    try:
        # Windowsの場合はantiword.exeが必要
        if platform.system() == 'Windows':
            # antiwordモジュールのパスを取得
            import antiword
            antiword_dir = os.path.dirname(antiword.__file__)
            
            # antiword.exeを探すか、antiword-winコマンドを実行する
            if shutil.which('antiword'):
                cmd = ['antiword', '-t', '-w', '0', doc_path]  # -t: テキスト出力, -w 0: 折り返しなし
            else:
                print("antiwordコマンドがインストールされていません。Pythonライブラリで代替します。")
                # 代替としてのpythonコードを実行（antiwordライブラリを使用）
                from antiword import process_file
                # 日本語対応のために適切なエンコーディングを指定
                text = process_file(doc_path)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                return output_path
        else:
            # Linux/Macの場合
            cmd = ['antiword', '-t', '-w', '0', doc_path]
        
        # antiwordコマンドを実行してテキストを抽出
        result = subprocess.run(cmd, capture_output=True, text=True, check=True, encoding='utf-8')
        
        # 結果をファイルに書き込む
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result.stdout)
        
        return output_path
    except Exception as e:
        # より詳細なエラー情報を出力
        print(f"antiwordでの変換に詳細なエラー: {str(e)}")
        raise e

def extract_text_with_custom_python(doc_path, output_path):
    """
    独自のPythonコードでdocファイルからテキストを抽出する試み
    日本語テキストの抽出に特化
    """
    try:
        # バイナリモードでファイルを開く
        with open(doc_path, 'rb') as f:
            content = f.read()
        
        # テキスト部分を抽出（ASCII文字と日本語文字）
        text_bytes = bytearray()
        
        # まず、明らかなXML/テキスト部分を探す
        xml_start = content.find(b'<?xml')
        if xml_start > 0:
            xml_content = content[xml_start:]
            # XMLをデコード
            try:
                xml_text = xml_content.decode('utf-8', errors='ignore')
                # XMLタグを除去
                xml_text = re.sub(r'<[^>]+>', ' ', xml_text)
                # 余分な空白を整理
                xml_text = re.sub(r'\s+', ' ', xml_text).strip()
                
                # XMLテキストを書き込む
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(xml_text)
                
                return output_path
            except Exception as xml_error:
                print(f"XML処理中のエラー: {str(xml_error)}")
        
        # ASCII文字と日本語文字（マルチバイト文字）を抽出
        i = 0
        while i < len(content):
            # ASCII文字と一部の制御文字
            if (32 <= content[i] <= 126) or content[i] in [9, 10, 13]:  # ASCII可読文字、タブ、改行
                text_bytes.append(content[i])
                i += 1
            # Shift-JIS/EUC-JPなどの日本語文字の可能性があるバイト
            elif (0x81 <= content[i] <= 0x9F or 0xE0 <= content[i] <= 0xEF) and i + 1 < len(content):
                # 2バイト文字として処理
                text_bytes.extend(content[i:i+2])
                i += 2
            else:
                i += 1
        
        # 複数のエンコーディングでデコードを試みる
        encodings = ['utf-8', 'shift_jis', 'euc-jp', 'iso-2022-jp', 'cp932']
        decoded_text = None
        
        for encoding in encodings:
            try:
                decoded_text = text_bytes.decode(encoding, errors='ignore')
                # 日本語文字が含まれているかチェック
                if re.search(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]', decoded_text):
                    break
            except UnicodeDecodeError:
                continue
        
        if decoded_text is None:
            decoded_text = text_bytes.decode('utf-8', errors='ignore')
        
        # テキストをクリーンアップ
        # 制御文字を削除
        decoded_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', decoded_text)
        # 連続する空白や制御文字を整理
        decoded_text = re.sub(r'\s+', ' ', decoded_text)
        # 意味のある文字列を検出（少なくとも1つの日本語文字または3文字以上の単語）
        meaningful_text = []
        for line in decoded_text.split('\n'):
            line = line.strip()
            if re.search(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]', line) or re.search(r'\w{3,}', line):
                meaningful_text.append(line)
        
        # 意味のあるテキストが見つかった場合
        if meaningful_text:
            final_text = '\n'.join(meaningful_text)
        else:
            final_text = decoded_text
        
        # テキストファイルに書き込む
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_text)
        
        return output_path
    except Exception as e:
        print(f"カスタムPython処理でのエラー: {str(e)}")
        raise e

def convert_doc_to_docx_then_text(doc_path, output_path):
    """
    .docファイルを一旦.docxに変換してからテキストに変換する
    """
    # 絶対パスに変換
    doc_path = os.path.abspath(doc_path)
    output_path = os.path.abspath(output_path)
    
    # 一時ファイルの作成
    temp_dir = tempfile.mkdtemp()
    temp_file = os.path.join(temp_dir, "temp.docx")
    
    try:
        # Wordアプリケーションの起動
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        
        try:
            # docファイルを開く
            print(f"Word COMでファイルを開いています: {doc_path}")
            doc = word.Documents.Open(doc_path, ReadOnly=True)
            
            # docxとして保存
            print(f"ファイルをDOCXとして保存中: {temp_file}")
            doc.SaveAs2(temp_file, FileFormat=16)  # 16 = docx
            
            # docxファイルを閉じる
            doc.Close()
            
            # docxをテキストに変換
            print("DOCXからテキストへの変換を実行中...")
            result = convert_docx_to_text(temp_file, output_path)
            
            return result
        
        except Exception as e:
            print(f"DOC→DOCX変換エラー（{doc_path}）: {str(e)}")
            traceback.print_exc()
            raise e
        
        finally:
            # Wordアプリケーションを終了
            try:
                word.Quit()
            except:
                pass
    
    except Exception as e:
        print(f"Word COM初期化エラー: {str(e)}")
        traceback.print_exc()
        raise e
    
    finally:
        # 一時ファイルの削除
        try:
            shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"一時ディレクトリの削除に失敗: {str(e)}")

def extract_text_with_powershell(doc_path, output_path):
    """
    PowerShellを使用してドキュメントからテキストを抽出する
    """
    # パスをエスケープ
    doc_path_escaped = doc_path.replace('\\', '\\\\').replace('"', '\\"')
    output_path_escaped = output_path.replace('\\', '\\\\').replace('"', '\\"')
    
    ps_command = f'''
    try {{
        $word = New-Object -ComObject Word.Application
        $word.Visible = $false
        $doc = $word.Documents.Open("{doc_path_escaped}")
        $text = $doc.Content.Text
        $doc.Close()
        $word.Quit()
        [System.Runtime.Interopservices.Marshal]::ReleaseComObject($word)
        $text | Out-File -FilePath "{output_path_escaped}" -Encoding utf8
        Write-Output "Success"
    }} catch {{
        Write-Error $_.Exception.Message
        exit 1
    }}
    '''
    
    # PowerShellコマンドを一時ファイルに保存
    ps_script = tempfile.NamedTemporaryFile(suffix=".ps1", delete=False)
    try:
        with open(ps_script.name, 'w', encoding='utf-8') as f:
            f.write(ps_command)
        
        # PowerShellスクリプトを実行
        print(f"PowerShellスクリプトを実行中: {ps_script.name}")
        process = subprocess.run(["powershell", "-ExecutionPolicy", "Bypass", "-File", ps_script.name], 
                      check=False, text=True, capture_output=True)
        
        if process.returncode != 0:
            print(f"PowerShellスクリプトのエラー出力: {process.stderr}")
            raise Exception(f"PowerShellスクリプトの実行に失敗: {process.stderr}")
        
        # 成功確認
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            print(f"PowerShellでのテキスト抽出に成功: {output_path}")
            return output_path
        else:
            raise Exception("出力ファイルが作成されませんでした")
            
    except Exception as e:
        print(f"PowerShellでのテキスト抽出エラー: {str(e)}")
        traceback.print_exc()
        raise e
    finally:
        # 一時スクリプトファイルを削除
        try:
            os.unlink(ps_script.name)
        except Exception as e:
            print(f"一時スクリプトファイルの削除に失敗: {str(e)}")

def process_directory(directory_path, recursive=True):
    """
    指定したディレクトリ内のすべてのWordファイル（.docと.docx）をテキストに変換する
    
    Args:
        directory_path (str): 処理するディレクトリのパス
        recursive (bool): サブディレクトリも再帰的に処理するかどうか
    
    Returns:
        tuple: (成功したファイルのリスト, 失敗したファイルのリスト)
    """
    # 絶対パスに変換
    directory_path = os.path.abspath(directory_path)
    print(f"ディレクトリを処理中: {directory_path}")
    
    # 成功・失敗したファイルのリスト
    success_files = []
    failed_files = []
    
    try:
        # 再帰的検索パターン
        if recursive:
            # Path.glob()で再帰的に検索
            docx_files = list(Path(directory_path).glob("**/*.docx"))
            doc_files = list(Path(directory_path).glob("**/*.doc"))
        else:
            # 非再帰
            docx_files = list(Path(directory_path).glob("*.docx"))
            doc_files = list(Path(directory_path).glob("*.doc"))
        
        print(f"検索結果: {len(docx_files)} DOCX ファイル, {len(doc_files)} DOC ファイル")
        
        # .docxファイルを処理
        for docx_file in docx_files:
            docx_file_str = str(docx_file)
            print(f"処理中: {docx_file_str}")
            try:
                output_path = convert_docx_to_text(docx_file_str)
                if output_path:
                    print(f"  変換完了: {output_path}")
                    success_files.append(docx_file_str)
                else:
                    print(f"  変換失敗: {docx_file_str}")
                    failed_files.append(docx_file_str)
            except Exception as e:
                print(f"  変換エラー（{docx_file_str}）: {str(e)}")
                traceback.print_exc()
                failed_files.append(docx_file_str)
        
        # .docファイルを処理
        for doc_file in doc_files:
            doc_file_str = str(doc_file)
            print(f"処理中: {doc_file_str}")
            try:
                output_path = convert_doc_to_text(doc_file_str)
                if output_path:
                    print(f"  変換完了: {output_path}")
                    success_files.append(doc_file_str)
                else:
                    print(f"  変換失敗: {doc_file_str}")
                    failed_files.append(doc_file_str)
            except Exception as e:
                print(f"  変換エラー（{doc_file_str}）: {str(e)}")
                traceback.print_exc()
                failed_files.append(doc_file_str)
    
    except Exception as e:
        print(f"ディレクトリ処理エラー: {str(e)}")
        traceback.print_exc()
    
    return success_files, failed_files

def extract_text_with_binary_parsing(doc_path, output_path):
    """
    バイナリデータから直接日本語テキストを抽出する
    複数のエンコーディングでテキストを抽出し、最も良質なものを選択する
    """
    try:
        print(f"バイナリ解析による日本語テキスト抽出を開始({doc_path})...")
        
        # バイナリモードでファイルを開く
        with open(doc_path, 'rb') as f:
            content = f.read()
        
        # 試すエンコーディングのリスト
        encodings = ['utf-8', 'utf-16le', 'utf-16be', 'shift_jis', 'euc-jp', 'cp932', 'iso-2022-jp']
        encoding_results = []
        
        # 各エンコーディングでの抽出を試みる
        for encoding in encodings:
            try:
                # ファイル全体をデコード
                text = content.decode(encoding, errors='ignore')
                
                # 日本語文字が含まれているか確認
                jp_chars = re.findall(r'[ぁ-んァ-ヶ一-龠々〆〜]', text)
                jp_ratio = len(jp_chars) / max(len(text), 1)
                
                print(f"  エンコーディング {encoding}: テキスト長={len(text)}, 日本語文字数={len(jp_chars)}, 比率={jp_ratio:.2%}")
                
                if len(text) > 100 and jp_ratio > 0.01:
                    # 意味のある段落だけを抽出
                    paragraphs = re.split(r'\r\n|\n|\r', text)
                    meaningful_paras = []
                    
                    for para in paragraphs:
                        # 日本語文字を含み、一定の長さがある段落のみ抽出
                        if re.search(r'[ぁ-んァ-ヶ一-龠々〆〜]', para) and len(para.strip()) > 5:
                            # 制御文字を削除
                            clean_para = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', para)
                            meaningful_paras.append(clean_para)
                    
                    if meaningful_paras:
                        filtered_text = '\n'.join(meaningful_paras)
                        # 再度日本語比率をチェック
                        jp_chars = re.findall(r'[ぁ-んァ-ヶ一-龠々〆〜]', filtered_text)
                        jp_ratio = len(jp_chars) / max(len(filtered_text), 1)
                        
                        if len(filtered_text) > 100 and jp_ratio > 0.05:
                            encoding_results.append((filtered_text, jp_ratio, encoding))
            except Exception as e:
                print(f"  エンコーディング {encoding} での抽出に失敗: {str(e)}")
        
        # バイナリデータから2バイト単位で日本語文字を直接抽出する試み
        try:
            i = 0
            chars = []
            current_chunk = []
            
            while i < len(content) - 1:
                # 2バイト単位で処理（UTF-16LE想定）
                if i + 2 <= len(content):
                    char_code = struct.unpack('<H', content[i:i+2])[0]
                    
                    # 日本語文字の範囲チェック
                    is_japanese = ((0x3040 <= char_code <= 0x309F) or  # ひらがな
                                  (0x30A0 <= char_code <= 0x30FF) or  # カタカナ
                                  (0x4E00 <= char_code <= 0x9FFF) or  # 漢字
                                  (char_code in [0x3000, 0x3001, 0x3002, 0xFF01, 0xFF0C, 0xFF0E, 0xFF1A, 0xFF1F]))  # 句読点など
                    
                    if is_japanese:
                        try:
                            char = content[i:i+2].decode('utf-16le', errors='ignore')
                            current_chunk.append(char)
                        except:
                            pass
                    else:
                        # 空白や改行なども取り込む
                        if char_code in [0x0020, 0x0009, 0x000A, 0x000D]:  # スペース、タブ、改行
                            try:
                                char = content[i:i+2].decode('utf-16le', errors='ignore')
                                current_chunk.append(char)
                            except:
                                pass
                        elif len(current_chunk) > 10:  # 十分な長さの日本語チャンクが見つかった
                            chars.append(''.join(current_chunk))
                            current_chunk = []
                i += 2
            
            # 最後のチャンクも追加
            if current_chunk and len(current_chunk) > 10:
                chars.append(''.join(current_chunk))
            
            if chars:
                # チャンク間の重複を削除
                unique_chunks = []
                for chunk in chars:
                    if chunk not in unique_chunks:
                        unique_chunks.append(chunk)
                
                binary_text = '\n'.join(unique_chunks)
                
                # 日本語比率を再チェック
                jp_chars = re.findall(r'[ぁ-んァ-ヶ一-龠々〆〜]', binary_text)
                jp_ratio = len(jp_chars) / max(len(binary_text), 1)
                
                if len(binary_text) > 100 and jp_ratio > 0.1:
                    encoding_results.append((binary_text, jp_ratio, "binary_direct"))
        except Exception as e:
            print(f"  バイナリ直接抽出に失敗: {str(e)}")
        
        # 結果を評価
        if encoding_results:
            # 日本語比率とテキスト長でソート
            encoding_results.sort(key=lambda x: (x[1], len(x[0])), reverse=True)
            best_text, best_ratio, best_encoding = encoding_results[0]
            
            print(f"最適なエンコーディング: {best_encoding} (日本語比率: {best_ratio:.2%}, 文字数: {len(best_text)})")
            
            # 余分な空行を整理
            best_text = re.sub(r'\n{3,}', '\n\n', best_text)
            
            # テキストファイルに書き込む
            with open(output_path, 'w', encoding='utf-8') as out_file:
                out_file.write(best_text)
            
            print(f"バイナリ解析によるテキスト抽出完了: {output_path}")
            return output_path
        else:
            raise Exception("有効な日本語テキストが見つかりませんでした")
    
    except Exception as e:
        print(f"バイナリ解析エラー: {str(e)}")
        traceback.print_exc()
        raise

def extract_text_doc_to_docx(doc_path, output_path):
    """
    .docファイルを一度.docxに変換してからテキストを抽出する
    """
    try:
        print(f"docからdocxへの変換を経由したテキスト抽出を開始({doc_path})...")
        
        # 一時的なdocxファイルパスを生成
        temp_docx_path = f"{os.path.splitext(doc_path)[0]}_temp.docx"
        
        # Windowsの場合はWord COMを使用
        if platform.system() == 'Windows':
            try:
                import win32com.client
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                
                try:
                    # docファイルを開く
                    doc = word_app.Documents.Open(os.path.abspath(doc_path), ReadOnly=True)
                    
                    # docxとして保存
                    doc.SaveAs2(os.path.abspath(temp_docx_path), FileFormat=16)  # 16はdocx形式
                    doc.Close(SaveChanges=False)
                finally:
                    word_app.Quit()
                
                # python-docxを使用してdocxからテキストを抽出
                import docx
                doc = docx.Document(temp_docx_path)
                
                # 段落を取得
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # テーブルからもテキストを抽出
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if para.text.strip():
                                    paragraphs.append(para.text)
                
                # 結果のテキストをファイルに書き込む
                text = '\n'.join(paragraphs)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                
                # 一時ファイルを削除
                if os.path.exists(temp_docx_path):
                    os.remove(temp_docx_path)
                
                return output_path
            except Exception as e:
                print(f"  Word COMでのdocx変換に失敗: {str(e)}")
                if os.path.exists(temp_docx_path):
                    os.remove(temp_docx_path)
                raise
        else:
            # Windowsでない場合はLibreOfficeを使用する（インストールされている必要がある）
            try:
                # LibreOfficeコマンドラインでの変換
                if platform.system() == 'Darwin':  # macOS
                    soffice_path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
                else:  # Linux
                    soffice_path = 'libreoffice'
                
                cmd = f'"{soffice_path}" --headless --convert-to docx --outdir "{os.path.dirname(temp_docx_path)}" "{doc_path}"'
                result = subprocess.run(cmd, shell=True, stderr=subprocess.PIPE)
                
                if result.returncode != 0:
                    error_message = result.stderr.decode('utf-8', errors='ignore')
                    raise Exception(f"LibreOffice変換エラー: {error_message}")
                
                # docxファイルからテキストを抽出
                import docx
                doc = docx.Document(temp_docx_path)
                
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # テーブルテキストも抽出
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if para.text.strip():
                                    paragraphs.append(para.text)
                
                # 結果のテキストをファイルに書き込む
                text = '\n'.join(paragraphs)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                
                # 一時ファイルを削除
                if os.path.exists(temp_docx_path):
                    os.remove(temp_docx_path)
                
                return output_path
            except Exception as e:
                print(f"  LibreOfficeでのdocx変換に失敗: {str(e)}")
                if os.path.exists(temp_docx_path):
                    os.remove(temp_docx_path)
                raise
                
        # どちらの方法も失敗した場合
        raise Exception("docxへの変換処理が失敗しました")
    
    except Exception as e:
        print(f"docからdocxへの変換経由でのテキスト抽出に失敗: {str(e)}")
        traceback.print_exc()
        raise

def main():
    if len(sys.argv) < 2:
        print("使用方法: python word_to_text_converter.py <マニュアル集のディレクトリパス> [--no-recursive]")
        return
    
    directory_path = sys.argv[1]
    recursive = True
    
    if len(sys.argv) > 2 and sys.argv[2] == "--no-recursive":
        recursive = False
    
    if not os.path.exists(directory_path):
        print(f"エラー: 指定されたパス '{directory_path}' が存在しません。")
        return
    
    if not os.path.isdir(directory_path):
        print(f"エラー: 指定されたパス '{directory_path}' はディレクトリではありません。")
        return
    
    print(f"ディレクトリ '{directory_path}' 内のWordファイルをテキストに変換します...")
    print(f"再帰的処理: {'有効' if recursive else '無効'}")
    
    success_files, failed_files = process_directory(directory_path, recursive)
    
    print("\n変換処理が完了しました。")
    print(f"成功: {len(success_files)}ファイル")
    print(f"失敗: {len(failed_files)}ファイル")
    
    if failed_files:
        print("\n失敗したファイル:")
        for file in failed_files:
            print(f"  - {file}")
        print("\n上記のファイルの変換に失敗しました。ファイルが開かれていないか、破損していないか確認してください。")

if __name__ == "__main__":
    main() 