#!/usr/bin/env python
# coding: utf-8

import os
import sys
import subprocess
import glob
from pathlib import Path
import docx2txt
import docx
import platform
import tempfile
import shutil
import re

def convert_docx_to_text(docx_path, output_path=None):
    """
    .docxファイルをテキストファイルに変換する
    """
    try:
        # 出力パスが指定されていない場合は入力ファイルと同じ場所に.txtを作成
        if output_path is None:
            output_path = str(Path(docx_path).with_suffix('.txt'))
        
        # docx2txtを使用してテキスト抽出を試みる
        try:
            text = docx2txt.process(docx_path)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"変換完了: {output_path}")
            return output_path
        except Exception as e1:
            print(f"docx2txtでの変換に失敗しました（{docx_path}）: {str(e1)}")
            print("python-docxでの変換を試みます...")
        
        # python-docxを使用して抽出を試みる
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # 空の段落を無視
                full_text.append(para.text)
        
        # テーブルからもテキストを抽出
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # 空のセルを無視
                        row_text.append(cell.text)
                if row_text:  # 空の行を無視
                    full_text.append('\t'.join(row_text))
        
        # テキストファイルに書き込む
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
        
        print(f"変換完了: {output_path}")
        return output_path
    
    except Exception as e:
        print(f"変換エラー（{docx_path}）: {str(e)}")
        return None

def convert_doc_to_text(doc_path, output_path=None, encoding='utf-8'):
    """
    .docファイルをテキストファイルに変換する
    Windowsの場合はCOMを使用して.docxに変換してから処理
    """
    try:
        # 出力パスが指定されていない場合は入力ファイルと同じ場所に.txtを作成
        if output_path is None:
            output_path = str(Path(doc_path).with_suffix('.txt'))
        
        # Windows環境の場合はWordのCOMを使用
        if platform.system() == 'Windows':
            try:
                import win32com.client
                
                # 一時ファイルの作成
                temp_dir = tempfile.mkdtemp()
                temp_docx = os.path.join(temp_dir, "temp.docx")
                
                try:
                    # Word COMを使用して.docを.docxに変換
                    print(f"Word COMを使用して変換中: {doc_path}")
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(doc_path))
                    doc.SaveAs2(os.path.abspath(temp_docx), FileFormat=16)  # 16 = .docx
                    doc.Close()
                    word.Quit()
                    
                    # 変換した.docxからテキストを抽出
                    print(f".docxからテキストを抽出中...")
                    try:
                        text = docx2txt.process(temp_docx)
                        
                        # 指定されたエンコーディングで書き込み
                        with open(output_path, 'w', encoding=encoding) as f:
                            f.write(text)
                        
                        print(f"変換完了: {output_path}")
                        return output_path
                    except Exception as e:
                        print(f"テキスト抽出に失敗: {str(e)}")
                        raise
                finally:
                    # 一時ディレクトリの削除
                    try:
                        shutil.rmtree(temp_dir)
                    except:
                        pass
            except ImportError:
                print("win32com.clientが利用できません。Python-docxでの変換を試みます...")
                
            except Exception as e:
                print(f"Word COMでの変換に失敗: {str(e)}")
                print("Python-docxでの変換を試みます...")
        
        # docx2txtを使ってdocファイルから直接テキストを抽出してみる
        try:
            print(f"docx2txtでdocファイルを直接変換中: {doc_path}")
            text = docx2txt.process(doc_path)
            
            # 指定されたエンコーディングで書き込み
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(text)
            
            print(f"変換完了: {output_path}")
            return output_path
        except Exception as e2:
            print(f"docx2txtでの直接変換に失敗: {str(e2)}")
        
        # バイナリ解析での変換を試みる
        try:
            print(f"バイナリ解析での変換を試みます: {doc_path}")
            
            # バイナリモードでファイルを開く
            with open(doc_path, 'rb') as f:
                content = f.read()
            
            # 複数のエンコーディングで試す
            encodings = ['utf-8', 'shift_jis', 'cp932', 'euc_jp']
            best_text = None
            best_jp_ratio = 0
            
            for enc in encodings:
                try:
                    # デコードを試みる
                    text = content.decode(enc, errors='ignore')
                    
                    # 日本語文字の割合を計算
                    jp_chars = len([c for c in text if '\u3040' <= c <= '\u30ff' or '\u4e00' <= c <= '\u9fff'])
                    jp_ratio = jp_chars / max(len(text), 1)
                    
                    if jp_ratio > best_jp_ratio:
                        best_text = text
                        best_jp_ratio = jp_ratio
                except:
                    continue
            
            if best_text and best_jp_ratio > 0.01:  # 1%以上の日本語文字があれば有効とする
                # 不要なバイナリノイズを除去
                cleaned_text = re.sub(r'[^\x20-\x7E\u3000-\u30FF\u4E00-\u9FFF\u3040-\u309F\uFF00-\uFF9F\u2000-\u206F\n]+', '', best_text)
                
                # 指定されたエンコーディングで書き込み
                with open(output_path, 'w', encoding=encoding) as f:
                    f.write(cleaned_text)
                
                print(f"バイナリ解析による変換完了: {output_path}")
                return output_path
            else:
                print("有効なテキストが抽出できませんでした")
        except Exception as e3:
            print(f"バイナリ解析での変換に失敗: {str(e3)}")
        
        print(f"すべての変換方法が失敗しました: {doc_path}")
        return None
        
    except Exception as e:
        print(f"変換エラー（{doc_path}）: {str(e)}")
        return None

def process_directory(directory_path, recursive=True, encoding='utf-8'):
    """
    指定したディレクトリ内のすべてのWordファイルをテキストに変換する
    """
    # 絶対パスに変換
    directory_path = os.path.abspath(directory_path)
    print(f"ディレクトリ処理中: {directory_path}")
    
    # 成功・失敗したファイルのリスト
    success_files = []
    failed_files = []
    
    # 再帰的検索パターン
    pattern = '**/*' if recursive else '*'
    
    # ファイルを検索
    docx_files = list(Path(directory_path).glob(f"{pattern}.docx"))
    doc_files = list(Path(directory_path).glob(f"{pattern}.doc"))
    
    print(f"発見したファイル: {len(docx_files)} DOCX ファイル, {len(doc_files)} DOC ファイル")
    
    # .docxファイルを処理
    for docx_file in docx_files:
        docx_path = str(docx_file)
        print(f"処理中: {docx_path}")
        try:
            output_path = convert_docx_to_text(docx_path)
            if output_path:
                success_files.append(docx_path)
            else:
                failed_files.append(docx_path)
        except Exception as e:
            print(f"変換エラー（{docx_path}）: {str(e)}")
            failed_files.append(docx_path)
    
    # .docファイルを処理
    for doc_file in doc_files:
        doc_path = str(doc_file)
        print(f"処理中: {doc_path}")
        try:
            output_path = convert_doc_to_text(doc_path, encoding=encoding)
            if output_path:
                success_files.append(doc_path)
            else:
                failed_files.append(doc_path)
        except Exception as e:
            print(f"変換エラー（{doc_path}）: {str(e)}")
            failed_files.append(doc_path)
    
    return success_files, failed_files

def main():
    if len(sys.argv) < 2:
        print("使用方法: python doc_to_txt.py <ファイル or ディレクトリ> [--no-recursive] [--encoding=ENCODING]")
        print("  --no-recursive: サブディレクトリを再帰的に処理しない")
        print("  --encoding=ENCODING: 出力エンコーディング (デフォルト: utf-8, 例: --encoding=shift-jis)")
        return
    
    path = sys.argv[1]
    recursive = True
    encoding = 'utf-8'
    
    # コマンドライン引数を解析
    for arg in sys.argv[2:]:
        if arg == "--no-recursive":
            recursive = False
        elif arg.startswith("--encoding="):
            encoding = arg.split("=")[1]
    
    if not os.path.exists(path):
        print(f"エラー: 指定されたパス '{path}' が存在しません。")
        return
    
    # ディレクトリかファイルかを判定
    if os.path.isdir(path):
        print(f"ディレクトリ '{path}' 内のWordファイルをテキストに変換します...")
        print(f"再帰的処理: {'有効' if recursive else '無効'}")
        print(f"出力エンコーディング: {encoding}")
        
        success_files, failed_files = process_directory(path, recursive, encoding)
        
        print("\n変換処理が完了しました。")
        print(f"成功: {len(success_files)}ファイル")
        print(f"失敗: {len(failed_files)}ファイル")
        
        if failed_files:
            print("\n失敗したファイル:")
            for file in failed_files:
                print(f"  - {file}")
    else:
        # 単一ファイルの処理
        file_path = path
        print(f"ファイル '{file_path}' をテキストに変換します...")
        print(f"出力エンコーディング: {encoding}")
        
        if file_path.lower().endswith('.docx'):
            convert_docx_to_text(file_path)
        elif file_path.lower().endswith('.doc'):
            convert_doc_to_text(file_path, encoding=encoding)
        else:
            print(f"エラー: サポートされていないファイル形式です。'.doc'または'.docx'ファイルを指定してください。")

if __name__ == "__main__":
    main() 