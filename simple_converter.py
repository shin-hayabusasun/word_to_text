#!/usr/bin/env python
# coding: utf-8

import os
import sys
import docx2txt
import docx
from pathlib import Path
import win32com.client
import tempfile
import shutil

def convert_file(file_path, encoding='shift-jis'):
    """
    .doc/.docxファイルをテキストに変換する
    """
    file_path = os.path.abspath(file_path)
    output_path = os.path.splitext(file_path)[0] + ".txt"
    
    print(f"変換中: {file_path}")
    print(f"出力先: {output_path}")
    
    # 拡張子を取得
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx':
        # .docxファイルの処理
        try:
            text = docx2txt.process(file_path)
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(text)
            print(f"docx2txtで変換成功")
            return True
        except Exception as e:
            print(f"docx2txtでの変換に失敗: {str(e)}")
            
            try:
                # python-docxを使用
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text)
                        if row_text:
                            full_text.append('\t'.join(row_text))
                
                with open(output_path, 'w', encoding=encoding) as f:
                    f.write('\n'.join(full_text))
                
                print(f"python-docxで変換成功")
                return True
            except Exception as e2:
                print(f"python-docxでの変換に失敗: {str(e2)}")
                return False
                
    elif ext == '.doc':
        # .docファイルの処理
        temp_dir = tempfile.mkdtemp()
        temp_docx = os.path.join(temp_dir, "temp.docx")
        
        try:
            # Word COMを使用
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            try:
                print(f"Word COMでdocxに変換中...")
                doc = word.Documents.Open(file_path)
                doc.SaveAs2(os.path.abspath(temp_docx), FileFormat=16)  # 16 = docx形式
                doc.Close(SaveChanges=False)
                
                # 変換したdocxを処理
                text = docx2txt.process(temp_docx)
                
                with open(output_path, 'w', encoding=encoding) as f:
                    f.write(text)
                
                print(f"Word COM + docx2txtで変換成功")
                return True
            except Exception as e:
                print(f"Word COMでの変換に失敗: {str(e)}")
                return False
            finally:
                word.Quit()
                try:
                    shutil.rmtree(temp_dir)
                except:
                    pass
        except Exception as e:
            print(f"Word COMの初期化に失敗: {str(e)}")
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
            return False
    else:
        print(f"サポートされていないファイル形式: {ext}")
        return False

# メイン処理
if __name__ == "__main__":
    # ファイル一覧
    files = [
        "2012.10給付金査定基準33 引受基準緩和型先進医療特約.doc",
        "201908給付金補助資料（手順書）11支払・不払件数の計上方法 (1).doc",
        "約款解釈①（照会・回答）＜2008.4～2011.5＞.docx"
    ]
    
    encoding = 'shift-jis'
    if len(sys.argv) > 1 and sys.argv[1].startswith('--encoding='):
        encoding = sys.argv[1].split('=')[1]
    
    print(f"出力エンコーディング: {encoding}")
    
    # 一覧のファイルをすべて変換
    success_count = 0
    failed_count = 0
    
    for file_path in files:
        if os.path.exists(file_path):
            print(f"\n処理開始: {file_path}")
            if convert_file(file_path, encoding):
                success_count += 1
                print(f"変換成功: {file_path}")
            else:
                failed_count += 1
                print(f"変換失敗: {file_path}")
        else:
            print(f"ファイルが見つかりません: {file_path}")
            failed_count += 1
    
    print(f"\n処理完了: 成功={success_count}, 失敗={failed_count}") 